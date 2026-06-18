"""Build real, editable .docx reports from a structured spec (python-docx).

Spec:
    {
      "title": "Q3 Report",                 # required
      "subtitle": "Prepared for the board",
      "sections": [
        {"heading": "Summary",
         "paragraphs": ["...", "..."],
         "bullets": ["...", "..."]},
        {"heading": "Numbers",
         "table": {"headers": ["Metric", "Q2", "Q3"],
                   "rows": [["Revenue", "14", "19"], ["Churn", "3%", "2%"]]}},
      ]
    }
"""

from __future__ import annotations

from typing import Any


def build_report(spec: dict[str, Any], out_path: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:  # pragma: no cover - optional dep
        raise RuntimeError("python-docx not installed (pip install python-docx)") from exc

    if not spec.get("title"):
        raise ValueError("report spec requires a 'title'")

    doc = Document()
    doc.add_heading(spec["title"], level=0)
    if spec.get("subtitle"):
        p = doc.add_paragraph(spec["subtitle"])
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(12)

    for section in spec.get("sections", []):
        if section.get("heading"):
            doc.add_heading(section["heading"], level=1)
        for para in section.get("paragraphs", []):
            doc.add_paragraph(para)
        for bullet in section.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")
        table_spec = section.get("table")
        if table_spec:
            _add_table(doc, table_spec)

    doc.save(out_path)
    return out_path


def _add_table(doc, table_spec: dict) -> None:
    headers = table_spec.get("headers", [])
    rows = table_spec.get("rows", [])
    ncols = len(headers) or (len(rows[0]) if rows else 0)
    if ncols == 0:
        return
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Light Grid Accent 1"
    if headers:
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(h)
            for r in cells[i].paragraphs[0].runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row[:ncols]):
            cells[i].text = str(val)
